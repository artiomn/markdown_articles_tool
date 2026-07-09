"""
DOCX formatter.
"""

import docx

from io import BytesIO
from pathlib import Path
import re

from markdown import markdown
from bs4 import BeautifulSoup


class DOCXFormatter:
    """
    Writes lines, into the DOCX.
    """

    format = 'docx'
    _header_re = re.compile(r'h([1-6])', flags=re.IGNORECASE)

    @staticmethod
    def _parse_block(doc, block):
        for element in block:
            if element is None or element.name is None:
                continue

            header_res = DOCXFormatter._header_re.search(element.name)
            if header_res is not None:
                doc.add_heading(element.text, level=int(header_res[1]))
            elif 'img' == element.name:
                doc.add_picture(element['src'])
            elif 'code' == element.name:
                paragraph = doc.add_paragraph()
            elif 'p' == element.name:
                paragraph = doc.add_paragraph()

                for child in element.children:
                    if child.name == 'strong':
                        paragraph.add_run(child.text).bold = True
                    elif child.name == 'em':
                        paragraph.add_run(child.text).italic = True
                    else:
                        paragraph.add_run(child)

                DOCXFormatter._parse_block(doc, element.children)

            elif 'ul' == element.name:
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Bullet')
            elif 'ol' == element.name:
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Number')

    @staticmethod
    def write(lines, **kwags):
        article_base_path = article_out_path.parent if (article_out_path := kwags.get('article_out_path'))\
                                                       is not None else Path.cwd()
        # Convert the string directly and save it.
        doc = docx.Document()

        # Adding content to the Word Document.
        DOCXFormatter._parse_block(doc, BeautifulSoup(markdown(lines, output_format='html'), 'html.parser'))

        stream = BytesIO()
        doc.save(stream)

        stream.seek(0)
        return stream.getvalue()
